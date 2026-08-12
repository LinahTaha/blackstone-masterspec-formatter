#Streamlit wraps code into webpage

import streamlit as st
import config #file
from openai import OpenAI
from docx import Document
import json #label 
import os #API key 
import io #input/output-> saves the file to memory bec on browser 
from docx.shared import RGBColor

#webpage title 
st.title("Blackstone Engineering – AI Agent Document Formatter")
st.write("Upload a messy spec document and get a clean MasterSpec formatted Word doc back.")

#upload button 
template_file = st.file_uploader("Upload your template doc", type="docx")
messy_file = st.file_uploader("Upload the doc to reformat", type="docx")

if template_file is not None and messy_file is not None:
    
    #opens the sample docx, extracts all the text into one string

    # extract text from template doc
    template_doc = Document(template_file)
    template_paragraphs = []
    for par in template_doc.paragraphs:
        if par.text.strip() != "":
            template_paragraphs.append(par.text)
    template_text = "\n".join(template_paragraphs)

    # extract text from messy doc
    messy_doc=Document(messy_file)
    messy_paragraphs=[]
    for p in messy_doc.paragraphs:
        if p.text.strip() != "":
            messy_paragraphs.append(p.text)
    messy_text="\n".join(messy_paragraphs)

     # save it as a file named whatever the user typed
    os.makedirs("formats", exist_ok=True)#create folder for formats 
    format_name = st.text_input("Name this format (e.g. NJDOT, PANYNJ)") 
                

    

    #button: nothing runs until clicked 
    if st.button("Format Document"):
        #spinner
        with st.spinner("Formatting your document..."):
            
            #send to openAi call 1
            client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=16000, #extend api call for longer documents 
                messages=[
                    {"role": "system",
                     "content": """You are a document parser for civil engineering specs. 
                     Analyze the structure and formatting of this document in detail.
                     Return ONLY a JSON object describing the heading levels, numbering system, 
                     and how sections and subsections are organized. No extra text, just the JSON."""
                    },
                    {"role": "user", "content": template_text}  
                ]

            )  
            #error message 1 + cleanup
    

            raw = response.choices[0].message.content #calls response 
            raw = raw.strip()                        # remove leading/trailing whitespace
            raw = raw.replace("```json", "")         # remove markdown json tag
            raw = raw.replace("```", "")            # remove closing backticks
            raw = raw.strip()                        # clean up again

            format_data = json.loads(raw, strict=False)
            
            with open(f"formats/{format_name}.json", "w") as f: #w-write
                json.dump(format_data, f) 
                

            response2 = client.chat.completions.create(
                model="gpt-4o-mini",
                max_tokens=16000,
                messages=[
                    {"role": "system",
                     "content": f"""You are a document parser for civil engineering specs.
                    This document follows this format structure:
                    {format_data}
                    
                    Extract ALL content from the document and return ONLY this exact JSON, using exactly these key names:
                    {{
                      "title": "the main document title",
                      "sections": [
                        {{
                          "heading": "section heading text",
                          "body": "body text or empty string",
                          "subsections": [
                            {{
                              "heading": "subsection heading text",
                              "body": "body text or empty string",
                              "subsections": []
                            }}
                          ]
                        }}
                      ]
                    }}
                    
                    Use 'heading' and 'body' and 'subsections' exactly. No other key names.
                    Always include the full section number in every heading e.g. '604 DRAINAGE STRUCTURES' not 'DRAINAGE STRUCTURES'.
                    Always include the full title with prefix in the title field e.g. 'SECTION 600 – INCIDENTAL CONSTRUCTION' not 'INCIDENTAL CONSTRUCTION'.""" 
                    },
                    {"role": "user",
                     "content": messy_text} 
                ]
            )

           

            raw2 = response2.choices[0].message.content
            raw2 = raw2.strip()
            raw2 = raw2.replace("```json", "")
            raw2 = raw2.replace("```", "")
            raw2 = raw2.strip()

            data= json.loads(raw2, strict=False)
    

            #out/JSON input/start
            def add_sections(doc, sections, level=1):
                #add current section header at level 
                for section in sections: 
                    
                    #error3: try all possible key names the AI might use
                    heading = section.get('heading') or section.get('title') or section.get('section_title') or ""
                    h=doc.add_heading(heading, level=level)
                    
                    # force heading color to black
                    for run in h.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)

                    #body but if nothing dont '...'
                    if section.get('body', '')!= '':
                        doc.add_paragraph(section['body'], style='Normal')
                        #calls its self again but level+1 (smaller)
                    if 'subsections' in section and section['subsections']:
                        add_sections(doc, section['subsections'], level+1) #calls that goes deeper 
                         
                          
            new_doc=Document()
            #add header: error2 instead of new_doc.add_heading(data['title'], level=0)
            # fix error4 - also look inside sections[0]
            title = (data.get('title') or 
                     data.get('section_title') or 
                     data.get('name') or
                     (data.get('sections', [{}])[0].get('section_title')) or
                     "Document")
           

            # page break after title
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            t = new_doc.add_heading(title, level=0)
            for run in t.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            new_doc.add_page_break()
            
            #calling    (doc, section[JSON]): error 2- AI uses slightly different key names it won't crash
            sections = data.get('sections') or data.get('subsections') or data.get('content') or []
            add_sections(new_doc, sections)
            
            

            # save it to memory instead of a file path-> send it to the browser, not save it locally
            buffer = io.BytesIO()
            new_doc.save(buffer)
            buffer.seek(0)

            #dowload button with finished doc
        st.success("Done! Download your formatted document below.")
        st.download_button(
            label="Download CleanOutput.docx",
            data=buffer,
            file_name="CleanOutput.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

